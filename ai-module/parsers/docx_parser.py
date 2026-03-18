"""
DOCX Resume Parser
Extracts text from DOCX files using python-docx
"""
from docx import Document
from typing import Optional
import os


def parse_docx(file_path: str) -> Optional[str]:
    """
    Extract text from DOCX resume
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text or None if parsing fails
    """
    try:
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return None
            
        doc = Document(file_path)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)
        
        # Join all text with newlines
        text = "\n".join(text_parts)
        
        return text.strip() if text else None
        
    except Exception as e:
        print(f"Error parsing DOCX: {e}")
        return None


def is_valid_docx(file_path: str) -> bool:
    """
    Check if file is a valid DOCX
    
    Args:
        file_path: Path to the file
        
    Returns:
        True if valid DOCX, False otherwise
    """
    try:
        if not file_path.lower().endswith('.docx'):
            return False
            
        # Try to open as Document
        doc = Document(file_path)
        return True
    except:
        return False
